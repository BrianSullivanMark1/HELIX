"""Extract plain text from a document so it can be ingested into a knowledge base.

Plain-text/markup files are read by the caller; this module handles the RICH formats — PDF and Word —
via optional libraries (pypdf / python-docx) imported LAZILY, so a missing library degrades to "this
format isn't supported right now" instead of crashing the app. Everything is best-effort: a corrupt or
encrypted file yields an empty string, never an exception into the ingest loop.

Scanned PDFs get a second chance: any page whose text layer is near-empty is handed to services/ocr
(Windows-native, on-machine) and read from its PIXELS. The decision is per PAGE, not per document, so
a mixed file — typed contract, scanned signature pages — comes through whole. OCR is capped in pages
and wall-clock so a thousand-page scan can't stall a turn; what goes untranscribed is said out loud in
the returned text rather than silently missing.
"""
from __future__ import annotations

from pathlib import Path

from helix.logging_setup import get_logger

_LOG = get_logger("doc_extract")

PDF_EXT = ".pdf"
DOCX_EXT = ".docx"
RICH_EXT = frozenset({PDF_EXT, DOCX_EXT})

_SCAN_TEXT_MIN = 40    # a real text-layer page has hundreds of chars; below this it's likely a scan
_OCR_MAX_PAGES = 100   # most pages OCR'd per document (~0.2s each → tens of seconds worst case)
_OCR_BUDGET_S = 30.0   # wall-clock ceiling for one document's OCR pass


def is_rich_doc(path) -> bool:
    """True for a format this module extracts (PDF / Word) rather than plain text."""
    return Path(path).suffix.lower() in RICH_EXT


def extract(path) -> str:
    """The text content of a PDF or Word document, or "" if unsupported, unreadable, or the optional
    library isn't installed. Never raises."""
    suffix = Path(path).suffix.lower()
    try:
        if suffix == PDF_EXT:
            return _extract_pdf(path)
        if suffix == DOCX_EXT:
            return _extract_docx(path)
    except Exception as exc:  # noqa: BLE001 - a bad document must never break the ingest loop
        _LOG.warning("could not extract %s: %s", path, exc)
    return ""


def _extract_pdf(path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        _LOG.info("pypdf not installed — PDF ingestion unavailable")
        return ""
    reader = PdfReader(str(path))
    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")  # try an empty-password decrypt; bail quietly if it won't open
        except Exception:  # noqa: BLE001
            return ""
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    pages = _ocr_scanned_pages(path, pages)
    return "\n\n".join(p for p in pages if p).strip()


def _ocr_scanned_pages(path, pages: list[str]) -> list[str]:
    """Fill in near-empty pages from their pixels. Per-page: a typed page keeps its text layer; a page
    with (almost) none is rendered and read by services/ocr. Whichever version says more wins — a
    cover page whose text layer is one heading loses nothing by being looked at twice. Appends a plain
    note when scans were left untranscribed — by the page/time caps, by a single page erroring, or by
    OCR failing outright on a document whose other pages were typed. The note names a cap only when
    the cap can be shown to have applied; otherwise it says what is missing and claims no cause."""
    scan_idxs = [i for i, t in enumerate(pages) if len(t) < _SCAN_TEXT_MIN]
    if not scan_idxs:
        return pages
    from helix.services import ocr  # lazy — pulls WinRT/pdfium only when a scan is actually present

    if not ocr.available():
        return pages
    read = ocr.scan_pdf_pages(path, scan_idxs, max_pages=_OCR_MAX_PAGES, budget_s=_OCR_BUDGET_S)
    if read:
        _LOG.info("OCR read %d scanned page(s) of %s", len(read), Path(path).name)
    out = list(pages)
    for i, text in read.items():
        if len(text) > len(out[i]):
            out[i] = text
    missed = [i for i in scan_idxs if i not in read]
    if missed and read:  # something bit into a real scan — say so instead of silently dropping pages
        # Only ONE of the three reasons a page can be absent from `read` is knowable from out here.
        # The page cap is arithmetic: scan_pdf_pages never looks past the first _OCR_MAX_PAGES
        # indices, so if we handed it more suspected scans than that, the cap demonstrably bit. The
        # wall-clock budget and a page that simply ERRORED mid-render (ocr._scan guards every page and
        # leaves the bad index out) are indistinguishable in the returned dict — and since the
        # per-page guard landed, the errored case is reachable for the first time. Naming "OCR
        # page/time limit" for it would tell the user, and the model reading this text, a cause that
        # is flatly wrong: nothing was rationed, a page was unreadable. So the cap is named only when
        # it provably applied, and everything else says what is lost without inventing a why.
        if len(scan_idxs) > _OCR_MAX_PAGES:
            out.append(f"… ({len(missed)} more scanned page(s) not transcribed — this document is "
                       "past the page limit for one OCR pass)")
        else:
            out.append(f"… ({len(missed)} more scanned page(s) not transcribed — their contents are "
                       "missing from this text)")
    elif missed:
        # Nothing came back at all, and we are past the ocr.available() gate, so OCR was there and
        # FAILED — a document pdfium won't open, or every page erroring in turn. The `and read` above is
        # a deliberate corroboration guard and stays: scan_idxs is only a SUSPICION (any page under
        # _SCAN_TEXT_MIN chars), so it routinely catches short TYPED pages — a cover sheet, a one-line
        # continuation page — and with no transcribed page to corroborate it we would libel a perfectly
        # clean typed PDF. A page whose text layer is COMPLETELY empty needs no corroboration, though:
        # typing produces characters, so zero characters means pixels. On a mixed contract those are
        # exactly the signature and exhibit pages, and the model must never reason over a contract that
        # quietly lost them.
        blank = [i for i in missed if not pages[i]]
        if blank:
            out.append(
                f"… ({len(blank)} scanned page(s) could not be read — their contents are missing "
                "from this text)"
            )
    return out


def _extract_docx(path) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        _LOG.info("python-docx not installed — Word ingestion unavailable")
        return ""
    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    # Include table cell text too — a lot of Word content lives in tables.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts).strip()
